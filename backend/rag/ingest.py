import os
import sys
import json
from datetime import datetime
from typing import List
from pathlib import Path
from typing import List, Dict, Any
from langchain_community.document_loaders import PyPDFLoader, UnstructuredWordDocumentLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.schema import Document

# 添加项目根目录到Python路径
current_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.dirname(current_dir)  # backend目录
if project_root not in sys.path:
    sys.path.insert(0, project_root)

# 现在可以统一使用绝对导入
from rag.resources import get_vector_db
from rag.utils.video2audio import AudioConverter

# 导入OCR处理器
# from .ocr_processor import get_ocr_processor

# 尝试导入更多PDF解析器
try:
    from langchain_community.document_loaders import PyMuPDFLoader
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    print("PyMuPDF未安装，将使用默认PDF解析器")

try:
    from langchain_community.document_loaders import PDFPlumberLoader
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False
    print("PDFPlumber未安装，将使用默认PDF解析器")

# 尝试导入更多Word文档解析器
try:
    from langchain_community.document_loaders import Docx2txtLoader
    HAS_DOCX2TXT = True
except ImportError:
    HAS_DOCX2TXT = False
    print("docx2txt未安装，将使用默认Word解析器")

# 尝试导入旧版Word文档解析器
try:
    import docx
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import parse_xml
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False
    print("python-docx未安装，无法解析旧版.doc文件")
    
from funasr import AutoModel

DB_DIR = os.path.join(os.path.dirname(__file__), 'db')
os.makedirs(DB_DIR, exist_ok=True)
AUDIO_TEXT_DIR = os.path.join(DB_DIR, 'audio_text')
os.makedirs(AUDIO_TEXT_DIR, exist_ok=True)
vector_db = get_vector_db()

# 全局ASR模型（单例）
_asr_model = None

SUPPORTED_VIDEO = {'.mp4', '.avi', '.mov', '.mkv', '.flv', '.wmv'}
SUPPORTED_AUDIO = {'.mp3', '.wav', '.m4a', '.aac', '.ogg'}

def get_asr_model():
    """获取ASR模型单例"""
    global _asr_model
    if _asr_model is None:
        print("加载语音识别模型...")
        _asr_model = AutoModel(
            model="paraformer-zh",
            vad_model="fsmn-vad",
            punc_model="ct-punc",
            # 长音频处理参数
            vad_kwargs={"max_single_segment_time": 30000},  # 30秒切片
            device="cpu",  # 或"cuda:0，注意pytorch要安装cuda版本"
            disable_update=True
        )
    return _asr_model

def is_media_file(file_path: str) -> str:
    """检查文件类型，返回 'video', 'audio' 或 None"""
    ext = os.path.splitext(file_path)[1].lower()
    if ext in SUPPORTED_VIDEO:
        return 'video'
    if ext in SUPPORTED_AUDIO:
        return 'audio'
    return 'unknown'

def process_media_file(file_path: str) -> List[Dict[str, Any]]:
    """
    音视频处理
    """
    media_type = is_media_file(file_path)
    if not media_type:
        return []
    
    audio_path = None
    cleanup_temp_audio = False

    try:
        print(f"处理{media_type}文件: {os.path.basename(file_path)}")
        
        # 音频预处理
        audio_path = file_path
        if media_type == 'audio':
            from pydub import AudioSegment
            import tempfile
            
            # 如果不是WAV格式，自动转换
            if not file_path.lower().endswith('.wav'):
                print(f"检测到非WAV格式 ({os.path.splitext(file_path)[1]})，正在转换...")
                
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.wav')
                temp_path = temp_file.name
                temp_file.close()
                
                audio = AudioSegment.from_file(file_path)
                audio = audio.set_frame_rate(16000).set_channels(1)
                audio.export(temp_path, format='wav')
                
                audio_path = temp_path
                cleanup_temp_audio = True
                print(f"转换完成: {temp_path}")
        
        # 视频文件处理
        elif media_type == 'video':
            print("提取音频...")
            import tempfile
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.wav')
            temp_path = temp_file.name
            temp_file.close()
            audio_path = str(AudioConverter.extract_audio_from_video(Path(file_path), output_path=Path(temp_path)))
            cleanup_temp_audio = True
        
        # 语音识别
        print("语音识别中...")
        model = get_asr_model()
        results = model.generate(
            input=audio_path,
            )
        
        print(f"识别完成，获得 {len(results)} 个结果")
        print(results)
        
        return results
        
    except Exception as e:
        print(f"处理失败: {e}")
        import traceback
        traceback.print_exc()
        return []
    finally:
        if cleanup_temp_audio and audio_path and os.path.exists(audio_path):
            try:
                os.remove(audio_path)
                print("临时音频文件已清理")
            except Exception as cleanup_error:
                print(f"清理临时音频文件失败: {cleanup_error}")


def save_asr_sentences_to_json(records: List[Dict[str, Any]], source_file: str) -> str:

    source_stem = Path(source_file).stem
    output_path = os.path.join(AUDIO_TEXT_DIR, f'{source_stem}.json')

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(records, f, ensure_ascii=False, indent=2)

    print(f"音频识别结果已保存: {output_path}")
    return output_path

# def is_scanned_pdf(file_path: str) -> bool:
#     """检测是否为扫描版PDF"""
#     try:
#         import fitz  # PyMuPDF
        
#         doc = fitz.open(file_path)
#         scanned_pages = 0
#         total_pages = len(doc)
        
#         for page_num in range(min(total_pages, 3)):  # 检查前3页
#             page = doc.load_page(page_num)
            
#             # 尝试提取文本
#             text = page.get_text()
            
#             # 如果文本很少或为空，可能是扫描版
#             if len(text.strip()) < 50:  # 少于50个字符
#                 scanned_pages += 1
        
#         doc.close()
        
#         # 如果超过一半的页面都是扫描版，则认为是扫描版PDF
#         return scanned_pages >= min(2, total_pages // 2)
        
#     except Exception as e:
#         print(f"检测扫描版PDF失败: {str(e)}")
#         return False

def parse_doc_file(file_path):
    """解析旧版.doc文件"""
    try:
        if HAS_PYTHON_DOCX:
            try:
                print("尝试使用python-docx解析.doc文件...")
                doc = docx.Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                if text.strip():
                    print("成功使用python-docx解析.doc文件")
                    return [{"page_content": text, "metadata": {}}]
            except Exception as e:
                print(f"python-docx解析失败: {str(e)}")
    except Exception as e:
        print(f"解析.doc文件时出错: {str(e)}")
        return None

def add_page_numbers_to_word(file_path: str) -> str:
    """为Word文档添加页码，返回带页码的临时文件路径"""
    try:
        if not HAS_PYTHON_DOCX:
            print("python-docx未安装，无法添加页码")
            return file_path
            
        # 打开Word文档
        doc = docx.Document(file_path)
        
        # 为每个节添加页码到页脚
        for section in doc.sections:
            footer = section.footer
            if not footer.paragraphs:
                footer.add_paragraph()
            
            paragraph = footer.paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 清除现有内容
            paragraph.clear()
            
            # 添加页码字段
            run = paragraph.add_run()
            fldSimple = parse_xml(r'<w:fldSimple xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:instr="PAGE" />')
            run._r.append(fldSimple)
        
        # 保存到临时文件
        import tempfile
        temp_dir = os.path.join(os.path.dirname(file_path), 'temp')
        os.makedirs(temp_dir, exist_ok=True)
        
        temp_file = os.path.join(temp_dir, f"temp_{os.path.basename(file_path)}")
        doc.save(temp_file)
        
        print(f"已为Word文档添加页码，临时文件：{temp_file}")
        return temp_file
        
    except Exception as e:
        print(f"添加页码失败: {str(e)}")
        return file_path

def process_word_with_pages(file_path: str) -> List[Document]:
    """处理Word文档并添加页码信息"""
    try:
        # 为Word文档添加页码
        temp_file = add_page_numbers_to_word(file_path)
        
        # 使用带页码的文档进行解析
        docs = None
        loaders_to_try = []
        
        # 按优先级排序Word解析器
        if HAS_DOCX2TXT:
            loaders_to_try.append(("Docx2txt", Docx2txtLoader))
        loaders_to_try.append(("UnstructuredWord", UnstructuredWordDocumentLoader))
        
        for loader_name, loader_class in loaders_to_try:
            try:
                print(f"尝试使用 {loader_name} 解析Word文档...")
                loader = loader_class(temp_file)
                docs = loader.load()
                if docs and any(len(doc.page_content.strip()) > 0 for doc in docs):
                    print(f"成功使用 {loader_name} 解析Word文档")
                    break
                else:
                    print(f"{loader_name} 解析结果为空，尝试下一个解析器")
            except Exception as e:
                print(f"{loader_name} 解析失败: {str(e)}")
                continue
        
        if not docs or not any(len(doc.page_content.strip()) > 0 for doc in docs):
            raise ValueError('所有Word解析器都无法提取到有效内容，可能是文件损坏或格式异常')
        
        # 为每个文档添加页码信息
        for i, doc in enumerate(docs):
            if not doc.metadata.get('page'):
                doc.metadata['page'] = i + 1
            doc.metadata['page_type'] = 'page'
        
        # 清理临时文件
        if temp_file != file_path and os.path.exists(temp_file):
            try:
                os.remove(temp_file)
                print(f"已清理临时文件：{temp_file}")
            except Exception as e:
                print(f"清理临时文件失败: {str(e)}")
        
        return docs
        
    except Exception as e:
        print(f"处理Word文档失败: {str(e)}")
        return []

def process_asr_result(asr_data: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """时间戳直接对应分词结果，直接返回JSON记录"""
    records = []
    
    for item in asr_data:
        text = item.get('text', '')
        text = text.replace(' ', '')
        timestamps = item.get('timestamp', [])
        
        # 分词
        import jieba
        words = list(jieba.cut(text))
        
        # 检查词数和时间戳数是否匹配
        if len(words) != len(timestamps):
            print(f"警告：词数({len(words)}) 与时间戳数({len(timestamps)}) 不匹配，将进行校正处理")
            # 如果时间戳不足，为多余的词添加假时间戳
            while len(timestamps) < len(words):
                if timestamps:
                    last_end = timestamps[-1][1]
                    timestamps.append([last_end, last_end + 100])
                else:
                    timestamps.append([0, 100])
        
        # 按标点分组
        current_words = []
        current_timestamps = []
        
        for i, word in enumerate(words):
            current_words.append(word)
            if i < len(timestamps):
                current_timestamps.append(timestamps[i])
            
            # 如果词以标点结尾，或这是最后一个词
            if word[-1] in "。！？；.!?;" or i == len(words) - 1:
                if current_words and current_timestamps:
                    sentence = "".join(current_words)
                    start_time = current_timestamps[0][0] / 1000.0
                    end_time = current_timestamps[-1][1] / 1000.0
                    
                    records.append({
                        'sentence': sentence,
                        'start_time': start_time,
                        'end_time': end_time
                    })
                    
                    current_words = []
                    current_timestamps = []
    return records


def asr_records_to_documents(records: List[Dict[str, Any]]) -> List[Document]:
    docs = []
    for item in records:
        docs.append(Document(
            page_content=item.get('sentence', ''),
            metadata={
                'start_time': item.get('start_time'),
                'end_time': item.get('end_time')
            }
        ))
    return docs

def ingest_file(file_path):
    media_type = is_media_file(file_path)
    if media_type:
        print(f"检测到媒体文件，类型: {media_type}")
        asr_result = process_media_file(file_path)
        
        if not asr_result:
            raise ValueError(f"媒体文件ASR处理返回空结果")
        
        print(f"ASR处理完成，获得 {len(asr_result)} 个结果")
        asr_records = process_asr_result(asr_result)
        
        if not asr_records:
            raise ValueError("ASR结果处理完成，但未生成任何Document")

        # docs = asr_records_to_documents(asr_records)
        
        filename = os.path.basename(file_path)
        file_path_str = str(file_path)  # 确保是字符串

        json_path = save_asr_sentences_to_json(asr_records, file_path_str)

        return json_path
    
    ext = os.path.splitext(file_path)[1].lower()
    # 尝试多种PDF解析器
    if ext == '.pdf':
        from paddleocr import PaddleOCR
        ocr = PaddleOCR(
            use_doc_orientation_classify=False,
            use_doc_unwarping=False,
            use_textline_orientation=False)

        # Run OCR inference on a sample image 
        result = ocr.predict(
            input="https://paddle-model-ecology.bj.bcebos.com/paddlex/imgs/demo_image/general_ocr_002.png")

        # Visualize the results and save the JSON results
        for res in result:
            res.print()
            res.save_to_img("output")
            res.save_to_json("output")
            
            

if __name__ == "__main__":
    test_file = "E:\\TraceLearn\\uploads\\asrtest.mp4"
    ingest_file(test_file)